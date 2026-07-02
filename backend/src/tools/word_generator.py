"""
Word Generation Tool - create_word

Creates Word documents with structured content.
Based on Anthropic's Agent Skills for docx document creation.
"""

import json
import base64
from typing import Annotated, Any
from langchain_core.tools import tool
from langchain_core.runnables import RunnableConfig

from ..runtime_context import user_id_ctx, thread_id_ctx
from ..sandbox import execute_python
from ..storage.file_storage import save_generated_file


def _validate_context(config: RunnableConfig | None) -> tuple[str, str]:
    """Validate and extract user_id and thread_id from config."""
    user_id = None
    thread_id = None
    
    if config:
        configurable = config.get("configurable") or {}
        if isinstance(configurable, dict):
            user_id = configurable.get("user_id")
            thread_id = configurable.get("thread_id")
    
    if not user_id:
        user_id = user_id_ctx.get()
    if not thread_id:
        thread_id = thread_id_ctx.get()
    
    if not user_id or user_id in ("default", "anonymous"):
        raise ValueError("Unable to identify the user. Please sign in again.")
    if not thread_id or thread_id == "default":
        raise ValueError("Unable to identify the thread. Please refresh the page.")
    
    return user_id, thread_id


@tool
def create_word(
    spec: Annotated[str, """
JSON specification for the Word document:
{
    "title": "Document title",
    "sections": [
        {"type": "heading", "level": 1, "text": "Chapter 1"},
        {"type": "paragraph", "text": "Body content..."},
        {"type": "bullet_list", "items": ["Point 1", "Point 2", "Point 3"]},
        {"type": "table", "headers": ["Column A", "Column B"], "rows": [["Data 1", "Data 2"]]}
    ]
}
    """],
    output_name: Annotated[str, "Output filename, such as 'document.docx'"],
    config: RunnableConfig = None,  # type: ignore[assignment]
) -> dict[str, Any]:
    """
    Create a Word document.
    
    Suitable for reports, contracts, memos, and other structured documents.
    
    Returns a structured object containing file_id, download_url, and summary.
    """
    try:
        user_id, thread_id = _validate_context(config)
    except ValueError as e:
        return {"success": False, "error": str(e)}
    
    if not output_name.endswith('.docx'):
        output_name += '.docx'
    
    try:
        spec_dict = json.loads(spec) if isinstance(spec, str) else spec
    except json.JSONDecodeError as e:
        return {"success": False, "error": f"Invalid JSON spec: {e}"}
    
    code = f'''
import json
import base64
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

spec = {json.dumps(spec_dict)}

doc = Document()

if spec.get("title"):
    title = doc.add_heading(spec["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

for section in spec.get("sections", []):
    section_type = section.get("type", "paragraph")
    
    if section_type == "heading":
        doc.add_heading(section.get("text", ""), level=section.get("level", 1))
    elif section_type == "paragraph":
        doc.add_paragraph(section.get("text", ""))
    elif section_type == "bullet_list":
        for item in section.get("items", []):
            doc.add_paragraph(item, style='List Bullet')
    elif section_type == "numbered_list":
        for item in section.get("items", []):
            doc.add_paragraph(item, style='List Number')
    elif section_type == "table":
        headers = section.get("headers", [])
        rows = section.get("rows", [])
        if headers:
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = str(header)
            for row_data in rows:
                row_cells = table.add_row().cells
                for i, value in enumerate(row_data):
                    if i < len(row_cells):
                        row_cells[i].text = str(value)

import io
buffer = io.BytesIO()
doc.save(buffer)
file_bytes = buffer.getvalue()

print("__FILE_BASE64_START__")
print(base64.b64encode(file_bytes).decode('utf-8'))
print("__FILE_BASE64_END__")
'''
    
    result = execute_python(code=code, timeout=60, user_id=user_id, thread_id=thread_id)
    
    if not result.get("success"):
        return {"success": False, "error": result.get("error") or result.get("stderr", "Execution failed")}
    
    stdout = result.get("stdout", "")
    try:
        start_idx = stdout.index("__FILE_BASE64_START__") + len("__FILE_BASE64_START__")
        end_idx = stdout.index("__FILE_BASE64_END__")
        file_bytes = base64.b64decode(stdout[start_idx:end_idx].strip())
    except Exception as e:
        return {"success": False, "error": f"Failed to retrieve file content: {e}"}
    
    file_info = save_generated_file(
        filename=output_name, data=file_bytes,
        user_id=user_id, thread_id=thread_id, file_type="generated"
    )
    
    if not file_info:
        return {"success": False, "error": "Failed to save file to storage"}
    
    sections = spec_dict.get("sections", [])
    
    return {
        "success": True,
        "file_id": file_info.get("file_id"),
        "download_url": file_info.get("download_url"),
        "filename": output_name,
        "file_type": "generated",
        "summary": f"Created Word document {output_name} with {len(sections)} sections",
    }


WORD_TOOLS = [create_word]
